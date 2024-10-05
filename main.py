import os
import pandas as pd
from docx import Document
from docx2pdf import convert  # To convert docx to pdf
from PyPDF2 import PdfMerger  # To merge PDFs
import re  # To sanitize filenames
import numpy as np  # To check for NaN values


# Load Excel file using pandas
excel_data = pd.read_excel('Park Road.xlsx')

# Template Word document path
template_doc_path = 'Membership-Information-Collection-Form-1.docx'

# Directory to save the generated documents
output_dir = 'generated_docs'
os.makedirs(output_dir, exist_ok=True)  # Create directory if it doesn't exist


def replace_placeholders(doc, placeholders):
    """
    Replaces placeholders in the Word document's paragraphs and tables.
    """
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if value is not None and not pd.isna(value):  # Only replace if value is valid
                if key in para.text:
                    para.text = para.text.replace(key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        if value is not None and not pd.isna(value):  # Only replace if value is valid
                            if key in para.text:
                                para.text = para.text.replace(key, value)


def sanitize_filename(filename):
    """
    Replaces any invalid characters in the filename with an underscore.
    """
    return re.sub(r'[\\/*?:"<>|]', '_', filename)


# Step 1: Generate Word documents and convert to PDFs
pdf_files = []  # To store paths of generated PDFs
for index, row in excel_data.iterrows():
    # Extract member data from the row, including the new fields
    category = row.get('CATEGORY', None)
    member_no = int(row.get('MEMBER_NO', 0))  # Ensure MEMBER_NO is an integer
    membership_date = row.get('MEMBERSHIP_DATE', None)
    current_status = row.get('CURRENT_STATUS', None)
    member_name = row.get('MEMBER_NAME', None)
    marital_status = row.get('MARITAL_STATUS', None)
    spouse_name = row.get('SPOUSE_NAME', None)
    father_name = row.get('FATHER_NAME', None)
    mothers_name = row.get('MOTHERS_NAME', None)
    date_of_marriage = row.get('DATE_OF_MARRIAGE', None)
    email = row.get('EMAIL', None)
    office_phone = row.get('Office_phone', None)
    address = row.get('ADDRESS', None)
    road_number = row.get('ROAD_NUMBER', None)
    house_number = row.get('HOUSE_NUMBER', None)
    blood_group = row.get('BLOOD_GROUP', None)
    gender = row.get('GENDER', None)
    mobile_number = row.get('MOBILE_NUMBER', None)
    apt_inside = row.get('apt_inside', None)
    address_outside = row.get('ADDRESS_OUTSIDE', None)
    roadnumber_outside = row.get('roadnumber_outside', None)
    house_number_outside = row.get('house_number_outside', None)
    apt_outside = row.get('apt_outside', None)

    # Load the template Word document
    doc = Document(template_doc_path)

    # Dictionary to map placeholders in the Word document to actual data, including the new fields
    placeholders = {
        '{CATEGORY}': str(category) if pd.notna(category) else '',
        '{MEMBER_NO}': str(member_no) if pd.notna(member_no) else '',
        '{MEMBERSHIP_DATE}': str(membership_date) if pd.notna(membership_date) else '',
        '{CURRENT_STATUS}': str(current_status) if pd.notna(current_status) else '',
        '{MEMBER_NAME}': str(member_name) if pd.notna(member_name) else '',
        '{MARITAL_STATUS}': str(marital_status) if pd.notna(marital_status) else '',
        '{SPOUSE_NAME}': str(spouse_name) if pd.notna(spouse_name) else '',
        '{FATHER_NAME}': str(father_name) if pd.notna(father_name) else '',
        '{MOTHERS_NAME}': str(mothers_name) if pd.notna(mothers_name) else '',
        '{DATE_OF_MARRIAGE}': str(date_of_marriage) if pd.notna(date_of_marriage) else '',
        '{EMAIL}': str(email) if pd.notna(email) else '',
        '{Office_phone}': str(office_phone) if pd.notna(office_phone) else '',
        '{ADDRESS}': str(address) if pd.notna(address) else '',
        '{ROAD_NUMBER}': str(road_number) if pd.notna(road_number) else '',
        '{HOUSE_NUMBER}': str(house_number) if pd.notna(house_number) else '',
        '{BLOOD_GROUP}': str(blood_group) if pd.notna(blood_group) else '',
        '{GENDER}': str(gender) if pd.notna(gender) else '',
        '{MOBILE_NUMBER}': str(mobile_number) if pd.notna(mobile_number) else '',
        '{apt_inside}': str(apt_inside) if pd.notna(apt_inside) else '',
        '{ADDRESS_OUTSIDE}': str(address_outside) if pd.notna(address_outside) else '',
        '{roadnumber_outside}': str(roadnumber_outside) if pd.notna(roadnumber_outside) else '',
        '{house_number_outside}': str(house_number_outside) if pd.notna(house_number_outside) else '',
        '{apt_outside}': str(apt_outside) if pd.notna(apt_outside) else '',
        # Category check marks
        '{A}': '✓' if category == 'AM' else '',
        '{G}': '✓' if category == 'GM' else '',
        '{L}': '✓' if category == 'LM' else '',
        '{S}': '✓' if category == 'SM' else '',
        '{D}': '✓' if category == 'DM' else '',
        '{T}': '✓' if category == 'TR-LM' else '',
        '{TR}': '✓' if category == 'TR-SM' else '',
        # Marital status check marks
        '{Married}': '✓' if marital_status == 'Married' else '',
        '{Single}': '✓' if marital_status == 'Single' else '',
        '{Divorced}': '✓' if marital_status == 'Divorced' else '',
        '{Separated}': '✓' if marital_status == 'Separated' else '',
        '{Widowed}': '✓' if marital_status == 'Widowed' else ''
    }

    # Replace placeholders with actual values
    replace_placeholders(doc, placeholders)

    # Sanitize the filename to prevent errors
    output_filename = sanitize_filename(f"{member_name.replace(' ', '_') if pd.notna(member_name) else 'Unknown_Member'}.docx")
    output_filepath = os.path.join(output_dir, output_filename)

    # Save the Word document
    doc.save(output_filepath)

    # Convert the Word document to PDF
    convert(output_filepath)  # This will generate a PDF in the same directory

    # Append the PDF path to the list
    pdf_files.append(output_filepath.replace('.docx', '.pdf'))

    print(f"Generated document for {member_name if pd.notna(member_name) else 'Unknown Member'}: {output_filepath.replace('.docx', '.pdf')}")

# Step 2: Merge all PDFs into a single file
merger = PdfMerger()
for pdf_file in pdf_files:
    merger.append(pdf_file)

# Output file path for the merged PDF
merged_pdf_path = os.path.join(output_dir, 'merged_documents.pdf')
merger.write(merged_pdf_path)
merger.close()

print(f"All documents have been merged into: {merged_pdf_path}")



#
# import os
# import pandas as pd
# from docx import Document
# from docx2pdf import convert  # To convert docx to pdf
# from PyPDF2 import PdfMerger  # To merge PDFs
# import re  # To sanitize filenames
# import numpy as np  # To check for NaN values
#
# # Load Excel file using pandas
# excel_data = pd.read_excel('Sproblem.xlsx')
#
# # Template Word document path
# template_doc_path = 'Bill Payee Information Collection Form.docx'
#
# # Directory to save the generated documents
# output_dir = 'generated_docs'
# os.makedirs(output_dir, exist_ok=True)  # Create directory if it doesn't exist
#
#
# def replace_placeholders(doc, placeholders):
#     """
#     Replaces placeholders in the Word document's paragraphs and tables.
#     """
#     # Replace placeholders in paragraphs
#     for para in doc.paragraphs:
#         for key, value in placeholders.items():
#             if value is not None and not pd.isna(value):  # Only replace if value is valid
#                 if key in para.text:
#                     para.text = para.text.replace(key, value)
#
#     # Replace placeholders in tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     for key, value in placeholders.items():
#                         if value is not None and not pd.isna(value):  # Only replace if value is valid
#                             if key in para.text:
#                                 para.text = para.text.replace(key, value)
#
#
# def sanitize_filename(filename):
#     """
#     Replaces any invalid characters in the filename with an underscore.
#     """
#     return re.sub(r'[\\/*?:"<>|]', '_', filename)
#
#
# # Step 1: Generate Word documents and convert to PDFs
# pdf_files = []  # To store paths of generated PDFs
# for index, row in excel_data.iterrows():
#     # Extract member data from the row
#     name = row.get('Name', None)
#     residence_id = row.get('Residence_ID', None)
#     road_number = row.get('Road_Number', None)
#     house_number = row.get('House_Number', None)
#     apt = row.get('APT', None)
#     category = row.get('Category', None)  # Category to set the right mark
#
#     # Load the template Word document
#     doc = Document(template_doc_path)
#
#     # Dictionary to map placeholders in the Word document to actual data
#     placeholders = {
#         '{name}': str(name) if pd.notna(name) else '',
#         '{Residence_ID}': str(residence_id) if pd.notna(residence_id) else '',
#         '{Road_Number}': str(road_number) if pd.notna(road_number) else '',
#         '{House_Number}': str(house_number) if pd.notna(house_number) else '',
#         '{APT}': str(apt) if pd.notna(apt) else '',
#         '{APARTMENT}': '✓' if category == 'APARTMENT' else '',
#         '{INDEPENDENT_HOUSE}': '✓' if category == 'INDEPENDENT HOUSE' else '',
#         '{ASSOCIATION_HOUSE}': '✓' if category == 'ASSOCIATION HOUSE' else '',
#         '{COMMERCIAL_FLAT}': '✓' if category == 'COMMERCIAL FLAT' else '',
#         '{EMBASSY}': '✓' if category == 'EMBASSY' else '',
#         '{UNDER_CONSTRUCTION}': '✓' if category == 'UNDER CONSTRUCTION' else ''
#     }
#
#     # Replace placeholders with actual values
#     replace_placeholders(doc, placeholders)
#     fname = str(residence_id)
#     # Sanitize the filename to prevent errors
#     output_filename = sanitize_filename(f"{fname.replace(' ', '_') if pd.notna(name) else 'Unknown_id'}.docx")
#     output_filepath = os.path.join(output_dir, output_filename)
#
#     # Save the Word document
#     doc.save(output_filepath)
#
#     # Convert the Word document to PDF
#     convert(output_filepath)  # This will generate a PDF in the same directory
#
#     # Append the PDF path to the list
#     pdf_files.append(output_filepath.replace('.docx', '.pdf'))
#
#     print(
#         f"Generated document for {name if pd.notna(name) else 'Unknown Name'}: {output_filepath.replace('.docx', '.pdf')}")
#
# # Step 2: Merge all PDFs into a single file
# merger = PdfMerger()
# for pdf_file in pdf_files:
#     merger.append(pdf_file)
#
# # Output file path for the merged PDF
# merged_pdf_path = os.path.join(output_dir, 'merged_documents.pdf')
# merger.write(merged_pdf_path)
# merger.close()
#
# print(f"All documents have been merged into: {merged_pdf_path}")